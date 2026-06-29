import streamlit as st
import openai
import os
import csv
import base64
import re
import requests
import urllib.parse
from datetime import datetime
from docx import Document
from docx.document import Document as DocumentType
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from PIL import Image
import io
import pdfplumber
import fitz  # PyMuPDF
from dotenv import load_dotenv

# Load environment variables
load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")

# Configuration
SUBMISSION_LOG = "submission_log.csv"
ADMIN_PASSWORD = os.getenv("ADMIN_PASSWORD", "admin123")


# ─────────────────────────────────────────────
# Submission log helpers
# ─────────────────────────────────────────────

def load_submissions():
    submissions = []
    if os.path.exists(SUBMISSION_LOG):
        try:
            with open(SUBMISSION_LOG, mode='r', newline='', encoding='utf-8') as file:
                reader = csv.DictReader(file)
                for row in reader:
                    submissions.append(row)
        except Exception as e:
            st.error(f"Error loading submissions: {e}")
    return submissions

def save_submissions(submissions):
    try:
        with open(SUBMISSION_LOG, mode='w', newline='', encoding='utf-8') as file:
            fieldnames = ["timestamp", "module", "groupnumber", "included_figures"]
            writer = csv.DictWriter(file, fieldnames=fieldnames)
            writer.writeheader()
            for submission in submissions:
                writer.writerow(submission)
        return True
    except Exception as e:
        st.error(f"Error saving submissions: {e}")
        return False

def log_submission(module, group_number, included_figures):
    submissions = load_submissions()
    new_submission = {
        "timestamp": datetime.now().isoformat(),
        "module": module,
        "groupnumber": str(group_number),
        "included_figures": str(included_figures)
    }
    submissions.append(new_submission)
    return save_submissions(submissions)

def get_submission_stats(submissions):
    if not submissions:
        return {"total": 0, "unique_groups": 0, "modules_with_submissions": 0}
    unique_groups = set()
    unique_modules = set()
    for submission in submissions:
        unique_groups.add(submission.get('groupnumber', ''))
        unique_modules.add(submission.get('module', ''))
    return {
        "total": len(submissions),
        "unique_groups": len(unique_groups),
        "modules_with_submissions": len(unique_modules)
    }

def get_submissions_by_module(submissions):
    modules = {}
    for submission in submissions:
        module = submission.get('module', '')
        if module not in modules:
            modules[module] = []
        modules[module].append(submission)
    return modules

def format_timestamp(timestamp_str):
    try:
        dt = datetime.fromisoformat(timestamp_str.replace('Z', '+00:00'))
        return dt.strftime('%Y-%m-%d %H:%M:%S')
    except:
        return timestamp_str


# ─────────────────────────────────────────────
# Google Docs helpers
# ─────────────────────────────────────────────

def extract_gdoc_id(url):
    """Extract the document ID from a Google Docs URL."""
    patterns = [
        r'/document/d/([a-zA-Z0-9_-]+)',
        r'id=([a-zA-Z0-9_-]+)',
    ]
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            return match.group(1)
    return None

def fetch_gdoc_as_docx(doc_id):
    """Download a Google Doc as a .docx file (bytes)."""
    export_url = f"https://docs.google.com/document/d/{doc_id}/export?format=docx"
    response = requests.get(export_url, timeout=30)
    if response.status_code == 200:
        return io.BytesIO(response.content)
    elif response.status_code == 403:
        raise PermissionError(
            "Could not access this Google Doc. Please make sure sharing is set to "
            "'Anyone with the link can view' before submitting."
        )
    else:
        raise RuntimeError(
            f"Failed to download Google Doc (HTTP {response.status_code}). "
            "Check that the link is correct and the document is shared publicly."
        )


# ─────────────────────────────────────────────
# PDF helpers
# ─────────────────────────────────────────────

def extract_text_from_pdf(pdf_bytes):
    """Extract all text from a PDF using pdfplumber."""
    text_parts = []
    with pdfplumber.open(pdf_bytes) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)

def extract_images_from_pdf(pdf_bytes):
    """Extract all images from a PDF using PyMuPDF (fitz)."""
    images = []
    pdf_bytes.seek(0)
    doc = fitz.open(stream=pdf_bytes.read(), filetype="pdf")
    for page_num in range(len(doc)):
        page = doc[page_num]
        for img_index, img in enumerate(page.get_images(full=True)):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            try:
                pil_image = Image.open(io.BytesIO(image_bytes))
                if pil_image.width > 100 and pil_image.height > 100:  # skip tiny icons
                    images.append(pil_image)
            except Exception:
                continue
    return images


# ─────────────────────────────────────────────
# docx helpers
# ─────────────────────────────────────────────

def extract_images_from_docx(doc):
    """Extract all images from a Word document."""
    images = []
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            try:
                image_data = rel.target_part.blob
                image = Image.open(io.BytesIO(image_data))
                images.append(image)
            except Exception as e:
                st.warning(f"Could not extract an image: {e}")
    return images


# ─────────────────────────────────────────────
# OpenAI vision helper
# ─────────────────────────────────────────────

def encode_image_for_api(image):
    """Convert PIL Image to base64 string for API."""
    buffer = io.BytesIO()
    if image.mode in ("RGBA", "P"):
        image = image.convert("RGB")
    image.save(buffer, format="JPEG")
    return base64.b64encode(buffer.getvalue()).decode()

def analyze_images_with_gpt4_vision(images, module):
    """Analyze images using GPT-4 Vision."""
    if not images:
        return "No figures found in the document."

    image_prompt_file_path = f"prompts/image_rubric_{module.split(' ')[0]}.txt"
    try:
        with open(image_prompt_file_path, 'r', encoding='utf-8') as f:
            prompt = f.read()
    except FileNotFoundError:
        try:
            with open("prompts/image_rubric_default.txt", 'r', encoding='utf-8') as f:
                prompt = f.read()
        except FileNotFoundError:
            prompt = "Analyze these figures and provide feedback on clarity, appropriateness, and professional presentation standards."

    messages = [
        {
            "role": "system",
            "content": f"You are an expert peer reviewer evaluating scientific figures. {prompt}"
        }
    ]

    for i, image in enumerate(images):
        base64_image = encode_image_for_api(image)
        messages.append({
            "role": "user",
            "content": [
                {"type": "text", "text": f"Figure {i+1} of {len(images)}:"},
                {
                    "type": "image_url",
                    "image_url": {
                        "url": f"data:image/jpeg;base64,{base64_image}",
                        "detail": "high"
                    }
                }
            ]
        })

    # Final instruction to ensure all figures are reviewed
    messages.append({
        "role": "user",
        "content": (
            f"There are {len(images)} figure(s) in total. You MUST provide a feedback block for every "
            f"single one of them — Figure 1 through Figure {len(images)}. Do not stop after the first figure. "
            "If a figure looks good and has no issues, say so explicitly and name the specific strengths."
        )
    })

    try:
        response = openai.chat.completions.create(
            model="gpt-4o",
            messages=messages,
            max_tokens=3000
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Error analyzing images: {e}"


# ─────────────────────────────────────────────
# Admin panel
# ─────────────────────────────────────────────

def admin_panel():
    """Admin interface for managing submissions."""
    st.header("🔧 Admin Panel")

    if 'admin_authenticated' not in st.session_state:
        st.session_state.admin_authenticated = False

    if not st.session_state.admin_authenticated:
        password = st.text_input("Enter admin password:", type="password")
        if st.button("Login"):
            if password == ADMIN_PASSWORD:
                st.session_state.admin_authenticated = True
                st.success("Admin access granted!")
                st.rerun()
            else:
                st.error("Invalid password!")
        return

    submissions = load_submissions()

    if not submissions:
        st.info("No submissions found.")
        if st.button("🚪 Logout", type="secondary"):
            st.session_state.admin_authenticated = False
            st.rerun()
        return

    st.subheader("📊 Submission Statistics")
    stats = get_submission_stats(submissions)
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Total Submissions", stats["total"])
    with col2:
        st.metric("Unique Groups", stats["unique_groups"])
    with col3:
        st.metric("Modules with Submissions", stats["modules_with_submissions"])

    st.subheader("📋 Submissions by Module")
    modules_data = get_submissions_by_module(submissions)
    for module in sorted(modules_data.keys()):
        module_submissions = modules_data[module]
        with st.expander(f"{module} ({len(module_submissions)} submissions)"):
            for i, submission in enumerate(module_submissions):
                col1, col2, col3 = st.columns([2, 2, 2])
                with col1:
                    st.write(f"**Group:** {submission.get('groupnumber', 'N/A')}")
                with col2:
                    st.write(f"**Time:** {format_timestamp(submission.get('timestamp', ''))}")
                with col3:
                    st.write(f"**Figures:** {submission.get('included_figures', 'N/A')}")
                if i < len(module_submissions) - 1:
                    st.divider()

    st.subheader("🛠️ Management Options")
    col1, col2 = st.columns(2)

    with col1:
        st.write("**Remove specific submission:**")
        if submissions:
            selection_options = []
            for i, submission in enumerate(submissions):
                display_text = f"Group {submission.get('groupnumber', 'N/A')} - {submission.get('module', 'N/A')} ({format_timestamp(submission.get('timestamp', ''))})"
                selection_options.append((i, display_text))

            selected_index = st.selectbox(
                "Select submission to remove:",
                options=[opt[0] for opt in selection_options],
                format_func=lambda x: next(opt[1] for opt in selection_options if opt[0] == x)
            )

            if st.button("🗑️ Remove Selected Submission", type="secondary"):
                if 'confirm_single_removal' not in st.session_state:
                    st.session_state.confirm_single_removal = False
                if not st.session_state.confirm_single_removal:
                    st.session_state.confirm_single_removal = True
                    st.warning("Click again to confirm removal.")
                    st.rerun()
                else:
                    updated_submissions = [sub for i, sub in enumerate(submissions) if i != selected_index]
                    if save_submissions(updated_submissions):
                        st.success("Submission removed successfully!")
                        st.session_state.confirm_single_removal = False
                        st.rerun()

    with col2:
        st.write("**Reset specific group/module:**")
        reset_group = st.text_input("Group number to reset:")
        reset_module = st.selectbox("Module to reset:", [
            "All modules",
            "2 - Research Questions",
            "3 - Study Design",
            "4 - Human Research Ethics",
            "5 - Presenting Results",
            "6 - Discussion Section"
        ])

        if st.button("🔄 Reset Group Submissions", type="secondary"):
            if not reset_group:
                st.warning("Please enter a group number.")
            else:
                if 'confirm_group_reset' not in st.session_state:
                    st.session_state.confirm_group_reset = False
                if not st.session_state.confirm_group_reset:
                    st.session_state.confirm_group_reset = True
                    if reset_module == "All modules":
                        st.warning(f"Click again to confirm removal of ALL submissions for Group {reset_group}.")
                    else:
                        st.warning(f"Click again to confirm removal of Group {reset_group}'s submission for {reset_module}.")
                    st.rerun()
                else:
                    if reset_module == "All modules":
                        updated_submissions = [sub for sub in submissions if sub.get('groupnumber', '') != reset_group]
                        success_msg = f"All submissions for Group {reset_group} removed!"
                    else:
                        updated_submissions = [sub for sub in submissions
                                             if not (sub.get('groupnumber', '') == reset_group and
                                                   sub.get('module', '') == reset_module)]
                        success_msg = f"Group {reset_group}'s submission for {reset_module} removed!"
                    if save_submissions(updated_submissions):
                        st.success(success_msg)
                        st.session_state.confirm_group_reset = False
                        st.rerun()

    st.subheader("⚠️ Danger Zone")
    with st.expander("Reset All Submissions", expanded=False):
        st.warning("This will delete ALL submission records. This action cannot be undone!")
        confirm_text = st.text_input("Type 'RESET ALL' to confirm:")
        if st.button("🚨 RESET ALL SUBMISSIONS", type="primary"):
            if confirm_text == "RESET ALL":
                if save_submissions([]):
                    st.success("All submissions have been reset!")
                    st.rerun()
            else:
                st.error("Please type 'RESET ALL' to confirm.")

    st.subheader("📥 Export Data")
    if submissions:
        output = io.StringIO()
        fieldnames = ["timestamp", "module", "groupnumber", "included_figures"]
        writer = csv.DictWriter(output, fieldnames=fieldnames)
        writer.writeheader()
        for submission in submissions:
            writer.writerow(submission)
        csv_data = output.getvalue()
        st.download_button(
            label="📥 Download Submission Data (CSV)",
            data=csv_data,
            file_name=f"submissions_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
            mime="text/csv"
        )

    if st.button("🚪 Logout", type="secondary"):
        st.session_state.admin_authenticated = False
        st.rerun()


# ─────────────────────────────────────────────
# Main app
# ─────────────────────────────────────────────

# ── Helper: read text (and optionally images) from any supported source ──
def read_document(file_obj=None, file_type="docx", analyze_figures=False, key_prefix=""):
    """Read text and images from a docx, pdf, or return None if nothing provided."""
    full_text = None
    images = []
    source_label = ""

    tab_docx, tab_pdf, tab_gdoc = st.tabs([
        "📄 Upload Word (.docx)",
        "📑 Upload PDF",
        "🔗 Google Docs link"
    ])

    with tab_docx:
        uploaded_docx = st.file_uploader(
            "Upload your .docx file. If you used Google Docs, use the Google Docs tab instead.",
            type="docx",
            key=f"{key_prefix}_docx"
        )
        if uploaded_docx:
            try:
                doc = Document(uploaded_docx)
                full_text = "\n".join([para.text for para in doc.paragraphs])
                if analyze_figures:
                    images = extract_images_from_docx(doc)
                source_label = "Word document"
            except Exception as e:
                st.error(f"Could not read Word file: {e}")

    with tab_pdf:
        st.info(
            "PDF support works best for text content. "
            "If your document contains figures (Module 5), a Word file or Google Doc will give more reliable figure analysis."
        )
        uploaded_pdf = st.file_uploader(
            "Upload your PDF file.",
            type="pdf",
            key=f"{key_prefix}_pdf"
        )
        if uploaded_pdf:
            try:
                pdf_bytes = io.BytesIO(uploaded_pdf.read())
                full_text = extract_text_from_pdf(pdf_bytes)
                if not full_text.strip():
                    st.error("No text could be extracted from this PDF. It may be a scanned image. Please upload a Word file instead.")
                    full_text = None
                else:
                    source_label = "PDF"
                    if analyze_figures:
                        pdf_bytes.seek(0)
                        images = extract_images_from_pdf(pdf_bytes)
            except Exception as e:
                st.error(f"Could not read PDF: {e}")

    with tab_gdoc:
        st.info(
            "Paste a Google Docs link below. Make sure sharing is set to "
            "**'Anyone with the link can view'** — otherwise the import will fail."
        )
        gdoc_url = st.text_input(
            "Google Docs URL",
            placeholder="https://docs.google.com/document/d/...",
            key=f"{key_prefix}_gdoc_url"
        )
        fetch_button = st.button("Import from Google Docs", key=f"{key_prefix}_gdoc_btn")

        if fetch_button and gdoc_url:
            doc_id = extract_gdoc_id(gdoc_url)
            if not doc_id:
                st.error("Could not find a valid document ID in that URL. Please check the link and try again.")
            else:
                try:
                    with st.spinner("Importing from Google Docs..."):
                        docx_bytes = fetch_gdoc_as_docx(doc_id)
                        doc = Document(docx_bytes)
                        full_text = "\n".join([para.text for para in doc.paragraphs])
                        if analyze_figures:
                            images = extract_images_from_docx(doc)
                        source_label = "Google Doc"
                        st.success("Google Doc imported successfully!")
                except PermissionError as e:
                    st.error(str(e))
                except Exception as e:
                    st.error(f"Could not import Google Doc: {e}")

    return full_text, images, source_label


def main_app():
    """Main application interface."""
    st.title("AI Peer Reviewer for BIOC32")
    st.markdown(
        "This AI Peer Reviewer will provide feedback to help you improve your submission "
        "before it is evaluated by the teaching assistants. It will not provide a grade for your submission."
    )

    # ── Module selection ──
    module = st.selectbox("Select Module", [
        "2 - Research Questions",
        "3 - Study Design",
        "4 - Human Research Ethics",
        "5 - Presenting Results",
        "6 - Discussion Section"
    ])

    analyze_figures = (module == "5 - Presenting Results")

    # ── Prior module config ──
    prior_module_map = {
        "3 - Study Design":          ("2 - Research Questions",   "Module 2 (Research Questions)"),
        "4 - Human Research Ethics": ("3 - Study Design",         "Module 3 (Study Design)"),
        "5 - Presenting Results":    ("3 - Study Design",         "Module 3 (Study Design)"),
        "6 - Discussion Section":    ("5 - Presenting Results",   "Module 5 (Presenting Results)"),
    }
    friendly_module_name = {
        "2 - Research Questions":    "Module 2 (Research Questions)",
        "3 - Study Design":          "Module 3 (Study Design)",
        "4 - Human Research Ethics": "Module 4 (Human Research Ethics)",
        "5 - Presenting Results":    "Module 5 (Presenting Results)",
        "6 - Discussion Section":    "Module 6 (Discussion Section)",
    }
    needs_prior = module in prior_module_map

    # ── Context-specific instructions ──
    if module == "3 - Study Design":
        st.info("Upload your approved Module 2 submission below, then upload your Module 3 submission. The reviewer will use your approved research question to assess whether your study design will answer it.")
    elif module == "4 - Human Research Ethics":
        st.info("Upload your approved Module 3 submission below, then upload your Module 4 submission. The reviewer will use your study design to evaluate your ethics review.")
    elif module == "5 - Presenting Results":
        st.info("Upload your approved Module 3 submission below, then upload your Module 5 submission. The reviewer will use your study design to assess whether your results and figures are appropriate. Make sure figures are embedded in your document.")
    elif module == "6 - Discussion Section":
        st.info("Upload your approved Module 5 submission below, then upload your Module 6 submission. The reviewer will use your results to evaluate how well your discussion interprets them.")

    # ── Prior module upload (required for modules 3–6) ──
    prior_text = None
    if needs_prior:
        _, prior_label = prior_module_map[module]
        st.markdown(f"### Step 1: Upload your approved {prior_label} submission")
        prior_text, _, _ = read_document(key_prefix="prior", analyze_figures=False)
        if prior_text is None:
            st.warning(f"⚠️ You must upload your approved {prior_label} submission before the reviewer can analyze your current module.")

    # ── Current module upload ──
    if needs_prior:
        st.markdown(f"### Step 2: Upload your {friendly_module_name[module]} submission")
    full_text, images, source_label = read_document(key_prefix="current", analyze_figures=analyze_figures)

    # ── Block if prior module missing ──
    if needs_prior and full_text and prior_text is None:
        st.error(f"❌ Please upload your approved {prior_module_map[module][1]} submission above before submitting.")
        st.stop()

    # ── Analysis ──
    if full_text:
        if not full_text.strip():
            st.warning("The document appears to be empty. Please check your file and try again.")
        else:
            # Combine prior + current text for the API
            if needs_prior and prior_text:
                _, prior_label = prior_module_map[module]
                combined_text = (
                    f"=== PREVIOUSLY APPROVED SUBMISSION: {prior_label} ===\n"
                    f"{prior_text}\n\n"
                    f"=== CURRENT SUBMISSION UNDER REVIEW: {module} ===\n"
                    f"{full_text}"
                )
            else:
                combined_text = full_text

            # ── Load rubric ──
            rubric_file_path = f"prompts/rubric_{module.split(' ')[0]}.txt"
            try:
                with open(rubric_file_path, 'r', encoding='utf-8') as f:
                    rubric_prompt = f.read()
            except FileNotFoundError:
                st.error("Rubric prompt file not found. Please check the prompts directory.")
                st.stop()

            # ── Module 5: three separate API calls for three distinct feedback sections ──
            if module == "5 - Presenting Results":
                stats_feedback = ""
                results_text_feedback = ""
                image_feedback = ""

                # ── Part 1: Statistical Analysis Assessment ──
                stats_rubric = (
                    "You are a peer reviewer for a third-year human physiology course. "
                    "Your task is ONLY to assess the statistical analysis used in this submission. "
                    "Produce a clearly labelled section titled \'## Part 1: Statistical Analysis Assessment\'. "
                    "You must: (1) identify which statistical test(s) were used; "
                    "(2) give an explicit verdict — APPROPRIATE or NOT APPROPRIATE — for each test; "
                    "(3) explain your reasoning considering data type, distribution, group structure (paired/unpaired, 2-group vs multi-group); "
                    "(4) if not appropriate, name the correct alternative and explain why in 1-2 sentences; "
                    "(5) if no test is mentioned, state this clearly as a major problem and direct students "
                    "to the Data Visualization and Analysis Tool on the Quercus page for this course. "
                    "Do not comment on writing style, figures, or anything other than the statistical approach."
                )
                try:
                    with st.spinner("Part 1 of 3: Assessing statistical analysis..."):
                        r1 = openai.chat.completions.create(
                            model="gpt-4-turbo",
                            messages=[
                                {"role": "system", "content": stats_rubric},
                                {"role": "user", "content": combined_text}
                            ]
                        )
                        stats_feedback = r1.choices[0].message.content
                except Exception as e:
                    stats_feedback = f"Statistical analysis assessment unavailable: {e}"

                # ── Part 2: Results Text Assessment ──
                results_rubric = (
                    "You are a peer reviewer for a third-year human physiology course. "
                    "Your task is ONLY to assess the written Results text (not figures, not stats methods). "
                    "Produce a clearly labelled section titled \'## Part 2: Results Text Assessment\'. "
                    "Evaluate whether the Results text: "
                    "(1) adequately guides the reader through the main findings in a logical order; "
                    "(2) describes trends and directions clearly (e.g., increased, decreased, no change) without repeating exact numeric values already shown in figures; "
                    "(3) uses correct statistical language — significant findings reported as \'significantly higher/lower (P = 0.xxx)\', "
                    "non-significant findings as \'no significant difference (P = 0.xxx)\'; "
                    "(4) avoids mechanistic interpretation (that belongs in the Discussion); "
                    "(5) references each figure at the appropriate point in the narrative. "
                    "For each issue: quote the relevant sentence, explain the problem, and provide a suggested rewrite. "
                    "If the results text is well-written, say so explicitly and identify what it does well. "
                    "Do not comment on figures or statistical test choice — only the prose."
                )
                try:
                    with st.spinner("Part 2 of 3: Assessing results text..."):
                        r2 = openai.chat.completions.create(
                            model="gpt-4-turbo",
                            messages=[
                                {"role": "system", "content": results_rubric},
                                {"role": "user", "content": combined_text}
                            ]
                        )
                        results_text_feedback = r2.choices[0].message.content
                except Exception as e:
                    results_text_feedback = f"Results text assessment unavailable: {e}"

                # ── Part 3: Figure Assessment ──
                try:
                    with st.spinner("Part 3 of 3: Assessing figures..."):
                        if images:
                            st.success(f"Found {len(images)} figure(s) in the document.")
                            image_feedback = analyze_images_with_gpt4_vision(images, module)
                        else:
                            image_feedback = (
                                "No figures were found in the document. "
                                "If you have figures, make sure they are properly embedded "
                                f"in your {source_label}."
                            )
                except Exception as e:
                    image_feedback = f"Figure assessment unavailable: {e}"

                # ── Log and display ──
                log_submission(module, "N/A", True)
                st.success("✅ Submission Successfully Reviewed. See Feedback Below.")
                st.subheader("Peer Review Feedback")

                with st.expander("## 📊 Part 1: Statistical Analysis Assessment", expanded=True):
                    st.write(stats_feedback)

                with st.expander("## 📝 Part 2: Results Text Assessment", expanded=True):
                    st.write(results_text_feedback)

                with st.expander("## 🖼️ Part 3: Figure Assessment", expanded=True):
                    st.write(image_feedback)

            else:
                # ── All other modules: single API call ──
                image_feedback = ""
                try:
                    if module == "2 - Research Questions":
                        with st.spinner("Analyzing content and searching recent literature — this may take up to 30 seconds..."):
                            response = openai.responses.create(
                                model="gpt-4o",
                                tools=[{"type": "web_search_preview"}],
                                instructions=rubric_prompt,
                                input=(
                                    f"{combined_text}\n\n"
                                    "IMPORTANT: Before providing feedback, search the web for recent "
                                    "peer-reviewed literature (2019–present) directly related to this "
                                    "research question. Use your search results to: (1) assess whether "
                                    "this question has already been answered, (2) provide 2–4 real, "
                                    "specific citations (with authors, journal, year, and DOI or URL) "
                                    "that students could read or cite, and (3) identify any factual "
                                    "errors in the background the students have written."
                                )
                            )
                            text_feedback = ""
                            for block in response.output:
                                if hasattr(block, "content"):
                                    for content_block in block.content:
                                        if hasattr(content_block, "text"):
                                            text_feedback += content_block.text
                            if not text_feedback:
                                text_feedback = "No feedback was generated. Please try again."
                    elif module == "3 - Study Design":
                        with st.spinner("Analyzing study design and searching for comparable studies — this may take up to 30 seconds..."):
                            response = openai.responses.create(
                                model="gpt-4o",
                                tools=[{"type": "web_search_preview"}],
                                instructions=rubric_prompt,
                                input=(
                                    f"{combined_text}\n\n"
                                    "IMPORTANT: Search the web for 2-3 real published studies that used "
                                    "a similar experimental design to the one proposed (similar intervention, "
                                    "population, or outcome measures). Cite each study fully (authors, journal, "
                                    "year, DOI) and explain specifically what the students can learn from it "
                                    "to improve their design."
                                )
                            )
                            text_feedback = ""
                            for block in response.output:
                                if hasattr(block, "content"):
                                    for content_block in block.content:
                                        if hasattr(content_block, "text"):
                                            text_feedback += content_block.text
                            if not text_feedback:
                                text_feedback = "No feedback was generated. Please try again."
                    elif module == "4 - Human Research Ethics":
                        with st.spinner("Analyzing ethics review and searching for supporting literature — this may take up to 30 seconds..."):
                            response = openai.responses.create(
                                model="gpt-4o",
                                tools=[{"type": "web_search_preview"}],
                                instructions=rubric_prompt,
                                input=(
                                    f"{combined_text}\n\n"
                                    "IMPORTANT: Where the students have proposed mitigations or monitoring thresholds for harms, search the web for peer-reviewed literature (2019-present) that supports or challenges those thresholds and protocols. Embed relevant citations (authors, journal, year, DOI) directly within the specific issues where they strengthen or correct the students' rationale. Also search for any clinical guidelines or published safety protocols relevant to the study population or intervention described."
                                )
                            )
                            text_feedback = ""
                            for block in response.output:
                                if hasattr(block, "content"):
                                    for content_block in block.content:
                                        if hasattr(content_block, "text"):
                                            text_feedback += content_block.text
                            if not text_feedback:
                                text_feedback = "No feedback was generated. Please try again."
                    elif module == "6 - Discussion Section":
                        with st.spinner("Analyzing discussion and searching for relevant literature — this may take up to 30 seconds..."):
                            response = openai.responses.create(
                                model="gpt-4o",
                                tools=[{"type": "web_search_preview"}],
                                instructions=rubric_prompt,
                                input=(
                                    f"{combined_text}\n\n"
                                    "IMPORTANT: Search the web for recent peer-reviewed literature (2019-present) relevant to the physiological mechanisms and findings discussed by the students. For each weakness you identify — particularly where mechanistic reasoning is shallow, a claim lacks support, or an interpretation could be strengthened — embed a real citation (authors, journal, year, DOI) that the students could use to deepen their discussion. Prioritise primary research articles and reviews that directly address the variables and population in the submission."
                                )
                            )
                            text_feedback = ""
                            for block in response.output:
                                if hasattr(block, "content"):
                                    for content_block in block.content:
                                        if hasattr(content_block, "text"):
                                            text_feedback += content_block.text
                            if not text_feedback:
                                text_feedback = "No feedback was generated. Please try again."
                    else:
                        with st.spinner("Analyzing content..."):
                            response = openai.chat.completions.create(
                                model="gpt-4-turbo",
                                messages=[
                                    {"role": "system", "content": rubric_prompt},
                                    {"role": "user", "content": combined_text}
                                ]
                            )
                            text_feedback = response.choices[0].message.content
                except Exception as e:
                    st.error(f"OpenAI API error: {e}")
                    st.stop()

                # ── Log and display ──
                log_submission(module, "N/A", False)
                st.success("✅ Submission Successfully Reviewed. See Feedback Below.")
                st.subheader("Peer Review Feedback")
                st.markdown("### 📝 Content Analysis")
                st.write(text_feedback)

    else:
        st.info("Please upload your document(s) above to receive feedback.")


# ─────────────────────────────────────────────
# Entry point
# ─────────────────────────────────────────────

def main():
    params = st.query_params
    if params.get("admin") == "true":
        admin_panel()
    else:
        main_app()

if __name__ == "__main__":
    main()
